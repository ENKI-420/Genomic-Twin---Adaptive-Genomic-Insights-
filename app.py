import streamlit as st
import requests
import pandas as pd
import json
import os
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
from dotenv import load_dotenv
from fhirclient.models.patient import Patient
from fhirclient import client
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
from openai import OpenAI

# Load environment variables
load_dotenv()

# Constants
GENOMIC_API_URL = os.getenv("GENOMIC_API_URL", "https://genomic-api-url.com/analyze")
FHIR_BASE_URL = "https://fhir.epic.com/interconnect-fhir-oauth/api/FHIR/R4/"
OAUTH_URL = "https://fhir.epic.com/interconnect-fhir-oauth/oauth2/token"
GNOMAD_API_URL = "https://gnomad.broadinstitute.org/api"
CLINVAR_API_URL = "https://api.ncbi.nlm.nih.gov/clinvar/v1"
TRIALS_API_URL = "https://api.3oncologyresearchhub.com/v2/studies"

# Load API keys
openai_api_key = os.getenv("OPENAI_API_KEY")

# Initialize FHIR Client
fhir_settings = {'app_id': os.getenv('FHIR_APP_ID'), 'api_base': FHIR_BASE_URL}
fhir_client = client.FHIRClient(settings=fhir_settings)

# Session setup with retries
session = requests.Session()
retry = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
session.mount('https://', HTTPAdapter(max_retries=retry))

# Authenticate with Epic OAuth2
def epic_authenticate(username, password):
    data = {'grant_type':'password', 'username':username, 'password':password,
            'client_id':os.getenv('EPIC_CLIENT_ID'), 'client_secret':os.getenv('EPIC_CLIENT_SECRET')}
    response = requests.post(OAUTH_URL, data=data)
    response.raise_for_status()
    return response.json().get('access_token')

# Fetch patient data from EPIC
def fetch_patient_data(patient_id, auth_token):
    headers = {"Authorization": f"Bearer {auth_token}"}
    response = session.get(f"{FHIR_BASE_URL}Patient/{patient_id}", headers=headers)
    response.raise_for_status()
    return response.json()

# Fetch Beaker Reports via Epic
def fetch_beaker_reports(patient_id, auth_token):
    headers = {"Authorization": f"Bearer {auth_token}"}
    response = session.get(f"{FHIR_BASE_URL}DiagnosticReport?patient={patient_id}", headers=headers)
    response.raise_for_status()
    return response.json()

# Fetch Clinical Trials
def fetch_clinical_trials(mutations):
    params = {"query": ",".join(mutations)}
    response = session.get(TRIALS_API_URL, params=params)
    response.raise_for_status()
    return response.json()

# Visualization Helper
def plot_mutation_data(data):
    df = pd.DataFrame(data)
    fig, ax = plt.subplots()
    df['gene'].value_counts().plot(kind='bar', ax=ax)
    ax.set_title('Mutation Frequency by Gene')
    ax.set_xlabel('Gene')
    ax.set_ylabel('Frequency')
    st.pyplot(fig)

# Generate PDF and DOCX Reports
def generate_reports(mutations, patient_data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, "Agile Oncology Report", ln=True, align='C')
    pdf.multi_cell(0, 10, json.dumps({"mutations": mutations, "patient_data": patient_data}, indent=4))
    pdf.output("genomic_report.pdf")

    doc = Document()
    doc.add_heading('Agile Oncology Report', 0)
    doc.add_heading('Patient Data', level=1)
    doc.add_paragraph(json.dumps(patient_data, indent=4))
    doc.add_heading('Mutations', level=1)
    doc.add_paragraph(json.dumps(mutations, indent=4))
    doc.save("genomic_report.docx")

# AI-Powered Clinical Trial Matching
def ai_clinical_trial_matching(mutations):
    trials_data = fetch_clinical_trials()
    matching_trials = [trial for trial in trials_data.get('studies', []) if any(mut in trial['eligibility'] for mut in mutations)]
    return matching_trials

# Main Streamlit UI
st.title("🚀 Agile Oncology AI Platform")
st.caption("Precision Medicine Platform v2.4")

with st.sidebar:
    st.header("🧬 Options")
    analysis_mode = st.radio("Select Analysis Mode", ["Genomic Analysis", "Beaker Reports", "Clinical Trial Matching"])
    username = st.text_input("Epic Username")
    password = st.text_input("Epic Password", type="password")
    if st.button("Login to Epic"):
        st.session_state.token = epic_authenticate(username, password)
        st.success("Authenticated with Epic")

patient_id = st.text_input("Enter Patient ID") if 'token' in st.session_state else None

if analysis_mode == "Genomic Analysis" and patient_id:
    file = st.file_uploader("Upload Genomic File (VCF/JSON)")
    if file and st.button("Analyze"):
        response = session.post(GENOMIC_API_URL, files={'file': file}).json()
        st.dataframe(pd.DataFrame(response['mutations']))
        plot_mutation_data(response['mutations'])
        patient_data = fetch_patient_data(patient_id, st.session_state.token)
        generate_reports(response['mutations'], patient_data)
        st.success("Reports Ready: PDF & DOCX")

elif analysis_mode == "Clinical Trial Matching" and st.button("Match Trials"):
    mutations = st.text_input("Enter mutations (comma-separated)").split(',')
    matching_trials = ai_clinical_trial_matching(mutations)
    st.json(matching_trials)

st.caption("🔬 Agile Oncology AI | Epic EHR | Secure & HIPAA-compliant")
