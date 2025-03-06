import requests
import pandas as pd
import matplotlib.pyplot as plt
import json
import os
from dotenv import load_dotenv
from openai import OpenAI
from fpdf import FPDF
from docx import Document
import streamlit as st

# Load environment variables
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Constants
GENOMIC_API_URL = os.getenv("GENOMIC_API_URL", "https://genomic-api-url.com/analyze")

# Analyze genomic data
def analyze_genomic_data(uploaded_file):
    try:
        if uploaded_file.name.endswith('.vcf'):
            df = pd.read_csv(uploaded_file, sep='\t', comment='#')
        elif uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        else:
            raise ValueError("Unsupported file format.")
        required_cols = ['Gene', 'Mutation', 'Variant_Classification', 'Allele_Frequency']
        for col in required_cols:
            if col not in df.columns:
                df[col] = "Unknown"
        return df
    except Exception as e:
        raise ValueError(f"Data analysis failed: {str(e)}")

# Visualization Helper
def plot_mutation_data(mutations_df):
    if not mutations_df.empty:
        plt.figure(figsize=(10, 6))
        mutations_df['Gene'].value_counts().plot(kind='bar', color='skyblue')
        plt.title('Mutation Frequency by Gene')
        plt.xlabel('Gene')
        plt.ylabel('Frequency')
        plt.tight_layout()
        st.pyplot(plt)
    else:
        st.warning("No mutations data available for plotting.")

# AI Interpretation
def ai_genomic_interpretation(mutations_df):
    client = OpenAI(api_key=openai_api_key)
    genomic_summary = mutations_df.to_dict(orient="records")
    prompt = f"Interpret the clinical significance and provide actionable insights for these genomic mutations: {genomic_summary}"
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# Generate Reports
def generate_reports(mutations_df, patient_data, insights, format_type="PDF"):
    content = json.dumps({"mutations": mutations_df.to_dict(orient='records'), "patient_data": patient_data, "AI Insights": insights}, indent=4)

    if format_type == "PDF":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, "Agile Oncology Genomic Report", ln=True, align='C')
        pdf.multi_cell(0, 10, content)
        pdf.output("genomic_report.pdf")
        return "genomic_report.pdf"

    elif format_type == "DOCX":
        doc = Document()
        doc.add_heading('Agile Oncology Genomic Report', 0)
        doc.add_heading('Patient Data', level=1)
        doc.add_paragraph(json.dumps(patient_data, indent=4))
        doc.add_heading('Mutation Analysis', level=1)
        doc.add_paragraph(content)
        doc_file = "genomic_report.docx"
        doc.save(doc_file)
        return doc_file

# Streamlit UI
st.title("🧬 Agile Oncology Genomic Analysis Module")

uploaded_file = st.file_uploader("📂 Upload Genomic Data (VCF/CSV)", type=['vcf', 'csv'])
report_format = st.selectbox("📄 Select Report Format", ["PDF", "DOCX"])

if uploaded_file and st.button("Analyze and Generate Report 🚀"):
    with st.spinner("Analyzing genomic data..."):
        try:
            mutations_df = analyze_genomic_data(uploaded_file)
            insights = ai_genomic_interpretation(mutations_df)
            patient_data = {"Patient ID": "Unknown"}  # Placeholder patient data; replace with real EPIC fetch

            st.dataframe(mutations_df)
            plot_mutation_data(mutations_df)
            report_file = generate_reports(mutations_df, patient_data, insights, report_format)

            with open(report_file, "rb") as f:
                st.download_button(label="Download Report", data=f, file_name=report_file)

            st.success("Report generated successfully!")
        
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
